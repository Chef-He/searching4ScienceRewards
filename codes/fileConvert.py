import io
import os
import tempfile
import pythoncom
import uuid
import time
import psutil
from win32com import client as wc
from docx import Document
from pdf2docx import parse

def kill_word_processes():
    for proc in psutil.process_iter():
        try:
            if proc.name().lower() == "winword.exe":
                proc.kill()
                time.sleep(0.3)
        except (psutil.NoSuchProcess, psutil.AccessDenied):
            pass

def doc_to_docx(doc_bytes: bytes) -> bytes | None:
    pythoncom.CoInitialize()
    tmp_doc_path = ""
    tmp_docx_path = ""
    word = None
    doc = None

    try:
        unique_id = uuid.uuid4().hex
        temp_dir = tempfile.gettempdir()
        
        tmp_doc_path = os.path.join(temp_dir, f"temp_{unique_id}.doc")
        tmp_docx_path = os.path.join(temp_dir, f"temp_{unique_id}.docx")

        with open(tmp_doc_path, "wb") as f:
            f.write(doc_bytes)

        word = wc.DispatchEx("Word.Application")  
        word.Visible = False  
        word.DisplayAlerts = False  
        word.AutomationSecurity = 1  

        max_retries = 2
        for attempt in range(max_retries + 1):
            try:
                doc = word.Documents.Open(
                    FileName=tmp_doc_path,
                    ReadOnly=True,
                    ConfirmConversions=False,
                    AddToRecentFiles=False
                )

                # 执行转换保存
                doc.SaveAs(
                    FileName=tmp_docx_path,
                    FileFormat=16,
                )
                break  # 成功则退出重试循环

            except Exception as e:
                if attempt == max_retries:
                    raise RuntimeError(f"转换失败（已重试{max_retries}次）") from e
                
                print(f"转换失败，正在重试 ({attempt+1}/{max_retries}): {str(e)}")
                kill_word_processes()
                time.sleep(1)

        time.sleep(0.5)
        if not os.path.exists(tmp_docx_path):
            raise FileNotFoundError("生成的docx文件不存在")

        with open(tmp_docx_path, "rb") as f:
            return f.read()

    except Exception as e:
        print(f"转换失败: {str(e)}")
        return None
    finally:
        try:
            if doc:
                doc.Close(SaveChanges=0)
                del doc
        except: pass

        try:
            if word:
                word.Quit()
                del word
        except: pass

        kill_word_processes()
        pythoncom.CoUninitialize()

        for path in [tmp_doc_path, tmp_docx_path]:
            if not path or not os.path.exists(path):
                continue
            for _ in range(3):
                try:
                    os.remove(path)
                    break
                except Exception as e:
                    print(f"删除临时文件失败（{path}）: {str(e)}")
                    time.sleep(0.5)

def pdf_to_docx(pdf_content: bytes) -> bytes:
    tmp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    tmp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    print(f"正在转换 {os.path.basename(tmp_pdf.name)}...")
    try:
        tmp_pdf.write(pdf_content)
        tmp_pdf.close()
        parse(tmp_pdf.name, tmp_docx.name)
        path = tmp_docx.name
        tmp_docx.close()
        with open(path, 'rb') as f:
            return f.read()
            
    finally:
        pythoncom.CoUninitialize()
        os.remove(tmp_pdf.name)
        os.remove(tmp_docx.name)

def parse_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ')
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                table_text.append(' | '.join(row_text))

        if table_text:
            full_text.append('\n表格内容:\n' + '\n'.join(table_text))
    result = '\n\n'.join(full_text)
    #with open("result.txt", "w") as f:
    #    f.write(result)
    if not result:
        print("No text parsed from docx")
    return result