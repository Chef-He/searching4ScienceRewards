import re
import os
import tempfile
import pythoncom
import chardet
from time import sleep
from win32com import client
from docx import Document
from typing import List, Dict, Tuple
from pdf2docx import parse
import requests

headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3'
}

def detect_encoding(content: bytes) -> str:
    result = chardet.detect(content)
    return result['encoding'] or 'gb18030'


def pdf_to_docx(pdf_content: bytes) -> bytes:
    tmp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    tmp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    print(f"正在转换 {os.path.basename(tmp_pdf.name)}...")
    try:
        # 写入临时PDF文件
        tmp_pdf.write(pdf_content)
        tmp_pdf.close()
        parse(tmp_pdf.name, tmp_docx.name)
        path = tmp_docx.name
        tmp_docx.close()
        # 读取转换结果
        with open(path, 'rb') as f:
            return f.read()
            
    finally:
        pythoncom.CoUninitialize()
        os.remove(tmp_pdf.name)
        os.remove(tmp_docx.name)

def doc_to_docx(doc_content: bytes) -> bytes:


def parse_word(content: bytes) -> tuple[str, list]:
    # 创建临时文件
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
        tmp_file.write(content)
    
    try:
        doc = Document(tmp_file.name)
        text = "\n".join(p.text for p in doc.paragraphs)
        tables = [
            [
                [cell.text.strip() for cell in row.cells] 
                for row in table.rows
            ]
            for table in doc.tables
        ]
        
        return text, tables

    except:
        pass

def extract_award_type(text: str) -> str:
    """
    提取文档中所有有效的奖项类别
    返回示例：['一等奖', '技术创新奖', '科技进步奖']
    """
    award_pattern = re.compile(
        r'(?:[\d一二三四五六七八九十]+[\.、）)])?'  # 匹配前置编号如"一."、"（二）"
        r'([^\s.，。！？；："“”‘’(){}\[\]]+奖)'  # 核心匹配规则
    )
    
    exclude_keywords = {"科学技术奖", "获奖名单"}
    
    for line in text.split('\n')[:]:
        for match in award_pattern.finditer(line):
            award = match.group(1)
            if len(award) >= 3 and not any(exclude in award for exclude in exclude_keywords):
                return award

    return "未知奖项类型"

def extract_year(text: str) -> list[int]:
    """
    从文本中提取所有符合 20xx年 格式的年份
    返回包含整型年份的列表（如 [2023, 2024]）
    """
    # 正则表达式模式（匹配 2000-2099 年的范围）
    pattern = r'(20\d{2})年'
    
    # 查找所有匹配项
    matches = re.findall(pattern, text)
    
    # 转换为整数并去重
    years = list({int(year) for year in matches})
    
    return years[0]

def extract_province(text: str) -> str:
    ans = ""
    try:
        ans = text[text.find("省") - 2: text.find("省") + 1]
    except:
        ans = text[text.find("市") - 2: text.find("市") + 1]
    print(ans)
    return ans

def map_table_header(header_row: List[str]) -> dict:
    """
    返回格式：{"name": 列索引, "unit": 列索引, "project": 列索引}
    """
    header_mapping = {}
    exclusion_keywords = ["提名", "号"]  # 需要排除的关键词
    
    for idx, header in enumerate(header_row):
        # 清洗表头文本
        clean_header = re.sub(r'\s+', '', header).lower()
        
        # 排除包含"提名"的列
        if any(ek in clean_header for ek in exclusion_keywords):
            continue
        
        # 定义关键词匹配规则
        matching_rules = {
            "name": ["姓名", "人", "者"],
            "unit": ["单位", "机构", "组织"],
            "project": [ "项目","课题","研究"],
            "award_type": ["奖"]  # 奖项类型字段
        }
        
        # 检查并记录有效列
        for field, keywords in matching_rules.items():
            if any(kw in clean_header for kw in keywords):
                if field not in header_mapping:
                    header_mapping[field] = idx
                    break 
    return header_mapping

def process_tables(text: str, tables: list) -> List[Dict]:
    """统一处理PDF/Word表格"""
    results = []
    award_year = extract_year(text)
    province = extract_province(text)
    # 分割文本段落
    sections = re.split(r'\n{2,}', text)
    section_ptr = 0
    
    award_types = []
    while(section_ptr < len(sections)):
        award_type = extract_award_type(sections[section_ptr])
        if(award_type != "未知奖项类型"):
            award_types.append(award_type)
        section_ptr += 1

    if (len(award_types) != len(tables)):
        print("award_types and tables not match")
        award_types = ["表格数不匹配"] * len(tables)
    index = 0
    for table in tables:
        if not table or len(table) < 2:
            continue
            
        # 处理表头
        header_mapping = map_table_header(table[0])
        award_level = ""
        # 处理数据行
        for row in table[1:]:
            if len(row) < len(table[0]):
                continue
            if row[0] == row[1]:
                if extract_award_type(row[0]) != "未知奖项类型":
                    award_level = " " + row[0]
            record = {
                "province": province,
                "year": award_year,
                "award_type": (row[header_mapping["award_type"]].strip() + award_level) if "award_type" in header_mapping else (award_types[index] + award_level),
                "name": row[header_mapping["name"]].strip().replace('\n', ' ') if "name" in header_mapping else "",
                "unit": row[header_mapping["unit"]].strip().replace('\n', ' ') if "unit" in header_mapping else "",
                "project": row[header_mapping["project"]].strip() if "project" in header_mapping else ""
            }
            if record["name"] == record["project"]:
                continue
            results.append(record)
        index += 1
        section_ptr += 1  # 移动到下一段落
    
    return results


def getContent(content_url: str) -> List[Dict[str, str]]:
   
    results = []
    
    try:
        response = requests.get(content_url, headers=headers, timeout=15)
        response.raise_for_status()
        content = response.content

        # 统一存储容器
        extracted_data = []

        # PDF处理分支
        if content_url.lower().endswith('.pdf'):
            try:
                print("检测到pdf, 开始转换")
                content = pdf_to_docx(content)
                print("成功将pdf转为docx文档")
                try:
                    text, tables = parse_word(content)
                    if tables:
                        print(f"发现 {len(tables)} 个表格")
                        results = process_tables(text, tables)
                    else:
                        print("无表格文档（暂不处理）")
                        return []
                except RuntimeError as e:
                    print(f"处理pdf文档时发生错误: {str(e)}")
            except Exception as e:
                print(f"失败:{str(e)}")
            
        elif content_url.lower().endswith(('.doc')):
            try:
                print("检测到doc, 开始转换")
                content = doc_to_docx(content)
                print("成功将doc转为docx文档")
                try:
                    text, tables = parse_word(content)
                    if tables:
                        print(f"发现 {len(tables)} 个表格")
                        results = process_tables(text, tables)
                    else:
                        print("无表格文档（暂不处理）")
                        return []
                except RuntimeError as e:
                    print(f"处理doc文档时发生错误: {str(e)}")
            except Exception as e:
                print(f"失败:{str(e)}")

        elif content_url.lower().endswith(('.docx')):
            try:
                text, tables = parse_word(content)
                if tables:
                    print(f"发现 {len(tables)} 个表格")
                    results = process_tables(text, tables)
                else:
                    print("无表格文档（暂不处理）")
                    return []
            except RuntimeError as e:
                print(f"处理Word文档时发生错误: {str(e)}")
        
        else:
            print("不支持的文件类型")
            return results  # 返回空列表

        # 转换数据结构
        for record in extracted_data:
            results.append({
                "source_url": content_url,
                "award_type": record.get("award_type", ""),
                "name": record.get("name", ""),
                "unit": record.get("unit", ""),
                "project": record.get("project", "")
            })

        # 打印摘要信息
        if results:
            print(f"从本文件提取到 {len(results)} 条有效记录")
        else:
            print("未提取到有效数据")

    except Exception as e:
        print(f"处理失败: {str(e)}")
    
    finally:
        return results  # 确保始终返回列表